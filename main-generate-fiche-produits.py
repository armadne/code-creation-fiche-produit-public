import openpyxl
import time
import requests
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches  # Permet de redimensionner l'image

from docx.shared import Pt  # Pour la taille du texte
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Pour aligner le texte
from docx.oxml import OxmlElement  # Pour styliser le texte

import os  # Importer os pour gérer les dossiers


# CREATION DE DOSSIER DANS LESQUELS ON VA METTRE LES IMAGES DES FICHES PRODUITS DANS LE DOSSIER "IMG_FICHE_PRODUIT_BOT"
# CREATION DE DOSSIER DANS LESQUELS ON VA METTRE LES FICHES PRODUITS GENERER PAR LE BOT DANS LE DOSSIER "FICHE_PRODUITS_BOT"

# Créer les dossiers s'ils n'existent pas déjà
os.makedirs("FICHE_PRODUITS_BOT", exist_ok=True)
os.makedirs("IMG_FICHE_PRODUIT_BOT", exist_ok=True)



# Charger le fichier Excel
path = "C:\\Users\\madan\\Desktop\\test-excel\\Tarif_Dometic_Hospitality.xlsx"
wb_obj = openpyxl.load_workbook(path)
sheet_obj = wb_obj.active
column_sku = sheet_obj['A']

# Liste des SKU
list_num_sku = [str(int(cell.value)) for cell in column_sku if cell.value and isinstance(cell.value, (int, float))]

# Fonction pour télécharger une image
def download_image(url, filename):
    try:
        response = requests.get(url, stream=True)
        if response.status_code == 200:
            with open(filename, "wb") as file:
                file.write(response.content)
            print(f"✅ Image téléchargée : {filename}")
            return filename
        else:
            print(f"❌ Erreur lors du téléchargement : {url}")
            return None
    except Exception as e:
        print(f"⚠️ Exception lors du téléchargement : {e}")
        return None

# Ouvrir Playwright
with sync_playwright() as p:
    browser = p.chromium.launch(headless=True, slow_mo=2000)
    page = browser.new_page() # OUVRE UN NOUVEL ONGLET SUR GOOGLE CHROME

    for sku in list_num_sku:  # Itérer sur chaque SKU
        print(f"\n🔎 Recherche SKU : {sku}")

        page.goto("https://www.dometic.com/fr-fr/outdoor") # ALLER SUR LE SITE DOMETIC

        # Accepter les cookies si visible
        try:
            cookie_popup = page.locator("#CookieReportsOverlay") # ACCEPTE LES COOKIES 
            if cookie_popup.is_visible():
                page.click(".wscrOk")
                print("✅ Cookies acceptés")
        except Exception: # SI IL N'Y A PAS DE COOKIES
            # ON AFFICHE EN TERMINAL
            print("ℹ️ Pas de pop-up de cookies.")

        # Recherche du produit
        page.get_by_role("button", name="Rechercher").click() # SELECTIONNE LA BARRE DE RECHERCHE DU SITE DOMETIC
        search_input = page.locator("input[placeholder='Rechercher']") #CLIQUE SUR LA BARRE DE RECHERCHE DU SITE DOMETIC
        search_input.fill(sku) # ECRIT DANS LA BARRE DE TRECHERCHE LE NUMERO SKU D'UN PRODUIT
        search_input.press("Enter") # APPUYER SUR ENTREE POUR LANCER LA RECHERCHE
        time.sleep(5) 

        try:
            
            document = Document() # CREATION D'UN FICHIER WORD
            
            
            
            # CLIQUER SUR LE PRODUIT 
            page.wait_for_selector("a[href*='/fr-fr/professional/solutions']", timeout=15000) # APRES LA RECHERCHE VIA LE NUMERO SKU, SELECTIONNE LE PRODUIT POUR AFFICHER LE PRODUIT
            product_link = page.locator("a[href*='/fr-fr/professional/solutions']").first 
            product_link.click() # CLIQUE SUR LE PRODUIT POUR AFFICHER LE PRODUIT
            time.sleep(5)
            page.wait_for_load_state("load") # ATTEND QUE LA PAGE QUI PRESENTE LE PRODUIT SOIT BIEN CHARGER
            
            

            # TITRE
            title_product = page.locator("h1").text_content().strip()
            print(f"📌 Produit trouvé : {title_product}")
            document.add_paragraph(f"Title: {title_product}")
            

            # IMAGE DU PRODUIT
            images = page.query_selector_all("img")
            image_urls = [img.get_attribute("src") for img in images if img.get_attribute("src") and img.get_attribute("src").startswith("http")]
            
            image_path = None
            if image_urls:
                # Télécharger l'image dans le dossier IMG_FICHE_PRODUIT_BOT
                image_path = os.path.join("IMG_FICHE_PRODUIT_BOT", f"image_{sku}.jpg")
                if image_urls:
                    download_image(image_urls[2], image_path)
                    document.add_picture(image_path, width=Inches(3))
                
                

            # DESCRIPTION DU PRODUIT
            container = page.locator("#container")
            paragraph = container.locator("p").nth(0).text_content().strip()
            print(f"📝 Description récupérée : {paragraph[:100]}...")
            document.add_paragraph(f"Description: {paragraph}")
            
  

            # SPÉCIFICATIONS
            page.get_by_role("button", name="Spécifications").click()
            time.sleep(2)
            
            
            # Sélectionner toutes les div dont l'id commence par 'specifications-'
            specifications_divs = page.locator("[id^='specifications-']").all()
            
                        
                        
            # Ignorer les premiers titres inutiles
            titles_to_ignore = {"Généralités", "Dimensions", "Puissance", "Performance", "Supplémentaire", "Logistique"}
            filtered_specs = []

            for spec in specifications_divs:
                text = spec.text_content().strip()
                if text and text not in titles_to_ignore:  # Ignorer les sections inutiles
                    filtered_specs.append(text)

            # Fonction pour formater proprement les spécifications
            def format_specification(text):
                formatted_lines = []
                current_line = ""

                words = text.split()  # Découper en mots pour restructurer la mise en page
                for word in words:
                    if word.isnumeric() or "[" in word or "SKU" in word:  # Si c'est une valeur ou une mesure
                        if current_line:  
                            formatted_lines.append(current_line.strip())  # Ajouter la ligne formatée
                        current_line = word  # Redémarrer une nouvelle ligne avec le mot actuel
                    else:
                        current_line += " " + word  # Ajouter le mot à la ligne actuelle

                if current_line:
                    formatted_lines.append(current_line.strip())  # Ajouter la dernière ligne

                return formatted_lines

            # Ajouter les spécifications au document Word
            for text in filtered_specs:
                formatted_specs = format_specification(text)
                for spec in formatted_specs:
                    paragraph = document.add_paragraph()
                    
                    if ":" in spec:  # Exemple: "Numéro SKU: 9600050533"
                        title, value = spec.split(":", 1)
                        run = paragraph.add_run(title.strip() + ": ")
                        run.bold = True  # Mettre en gras le titre
                        paragraph.add_run(value.strip())  # Ajouter la valeur
                    else:
                        # Si pas de séparation par ":", traiter normalement
                        paragraph.add_run(spec)

                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Ajouter une section pour "AVEC LE CARTON DU COLIS:"
            document.add_paragraph("\nAVEC LE CARTON DU COLIS (5 DERNIERES LIGNES):", style="Heading 2")
            

            time.sleep(3)       
                        
                
                
            # CLIQUE SUR "MANUEL ET DOCUMENTATION DU PRODUIT"
            page.get_by_role("button", name="Manuels et Documentation produit").click()
            time.sleep(2)
            
            
            # Récupère tous les liens dans la section Manuels et Documentation
            links = page.locator("div:has-text('Manuels') a").all()

            # Filtrer uniquement les liens PDF de documentation produit
            manual_links = [link.get_attribute('href') for link in links if link.get_attribute('href') and "externalassets" in link.get_attribute('href')]

            # Ajouter dans le document Word
            document.add_paragraph("📄 Manuels et Documentation produit:")
            for link in manual_links:
                document.add_paragraph(link)

            # Sauvegarde le document
            #document.save(f"fiche_produit_{sku}.docx")
            print(f"📄 Manuels ajoutés au document : {manual_links}")
                
                
            
            # Enregistrer le document Word dans le dossier FICHE_PRODUITS_BOT
            word_path = os.path.join("FICHE_PRODUITS_BOT", f"fiche_produit_{sku}.docx")
                

            # Enregistrer le document Word
            document.save(word_path)
            print(f"📄 Document Word enregistré : {word_path}")

        except Exception as e:
            print(f"❌ Une erreur s'est produite pour SKU {sku}: {e}")

    browser.close()  
    print("\n✅ Extraction terminée !")
